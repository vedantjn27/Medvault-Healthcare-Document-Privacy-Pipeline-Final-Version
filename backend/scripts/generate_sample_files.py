"""Generate synthetic, non-patient sample documents for every supported format."""

from __future__ import annotations

import io
import mailbox
import shutil
import sys
from datetime import datetime
from email.message import EmailMessage
from pathlib import Path

import fitz
import numpy as np
from docx import Document
from docx.shared import Inches
from openpyxl import Workbook
from openpyxl.drawing.image import Image as XlsxImage
from openpyxl.styles import Font, PatternFill
from PIL import Image, ImageDraw, ImageFont
from pydicom import Dataset, FileDataset, FileMetaDataset
from pydicom.uid import ExplicitVRLittleEndian, generate_uid


ROOT = Path(__file__).resolve().parents[1] / "sample_files"
SYNTHETIC = {
    "name": "Jordan Example", "email": "jordan.example@example.com",
    "phone": "415-555-0188", "mrn": "MV-482910", "npi": "1234567893",
    "dob": "1984-01-15", "condition": "diabetes mellitus",
    "medication": "metformin", "address": "42 Example Avenue, Boston, MA 02110",
}


def font(size: int):
    for path in (Path("C:/Windows/Fonts/arial.ttf"), Path("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf")):
        if path.is_file():
            return ImageFont.truetype(str(path), size)
    return ImageFont.load_default()


def phi_lines(extra: tuple[str, ...] = ()) -> list[str]:
    return [
        f"Patient: {SYNTHETIC['name']}", f"Email: {SYNTHETIC['email']}",
        f"Phone: {SYNTHETIC['phone']}", f"MRN: {SYNTHETIC['mrn']}",
        f"NPI: {SYNTHETIC['npi']}", f"DOB: {SYNTHETIC['dob']}",
        f"Address: {SYNTHETIC['address']}", f"Condition: {SYNTHETIC['condition']}",
        f"Medication: {SYNTHETIC['medication']}", *extra,
    ]


def text_image(title: str, lines: list[str], size=(1400, 1000)) -> Image.Image:
    image = Image.new("RGB", size, "white")
    draw = ImageDraw.Draw(image)
    draw.text((50, 45), title, fill="black", font=font(42))
    for index, line in enumerate(lines):
        draw.text((60, 130 + index * 78), line, fill="black", font=font(32))
    return image


def pdf(path: Path, title: str, lines: list[str], *, scanned_page: bool = False) -> None:
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((54, 52), title, fontsize=16)
    for index, line in enumerate(lines):
        page.insert_text((60, 95 + index * 38), line, fontsize=11)
    if scanned_page:
        scan = text_image("Scanned attachment", lines[:5], (1400, 900))
        buffer = io.BytesIO()
        scan.save(buffer, format="JPEG", quality=90)
        second = doc.new_page()
        second.insert_image(second.rect, stream=buffer.getvalue())
    doc.save(path)
    doc.close()


def create_pdfs() -> None:
    folder = ROOT / "pdf"
    folder.mkdir(parents=True)
    pdf(folder / "patient_portal_mode.pdf", "Patient Portal Mode", phi_lines(("Subject patient ID: MV-482910",)))
    pdf(folder / "research_sharing_mode.pdf", "Research Sharing Mode", phi_lines(("Age: 42", "Gender: female", "ZIP: 021")))
    pdf(folder / "insurance_processing_mode.pdf", "Insurance Processing Mode", phi_lines(("Procedure code: 99213", "Diagnosis code: E11.9", "Payer ID: 842610001")))
    pdf(folder / "legal_discovery_mode.pdf", "Legal Discovery Mode", phi_lines(("Privileged and confidential attorney-client legal advice.", "Outside counsel work product.")))
    pdf(folder / "custom_mode_mixed_native_scanned.pdf", "Custom Mode — Mixed Native and Scanned", phi_lines(), scanned_page=True)


def create_docx() -> None:
    folder = ROOT / "docx"
    folder.mkdir(parents=True)
    native = Document()
    native.add_heading("Synthetic Clinical Summary", 0)
    native.sections[0].header.paragraphs[0].text = f"MRN: {SYNTHETIC['mrn']}"
    table = native.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text, table.rows[0].cells[1].text = "Field", "Value"
    for line in phi_lines():
        key, value = line.split(":", 1)
        cells = table.add_row().cells
        cells[0].text, cells[1].text = key, value.strip()
    native.sections[0].footer.paragraphs[0].text = f"Contact: {SYNTHETIC['email']}"
    native.save(folder / "native_text_tables_headers.docx")

    mixed = Document()
    mixed.add_heading("Mixed Text and Embedded Image", 0)
    mixed.add_paragraph(f"Patient: {SYNTHETIC['name']} — Phone: {SYNTHETIC['phone']}")
    image_path = folder / "_embedded_source.png"
    text_image("Scanned Referral", phi_lines()[:6]).save(image_path)
    mixed.add_picture(str(image_path), width=Inches(6))
    mixed.save(folder / "mixed_text_and_embedded_image.docx")
    image_path.unlink()


def create_xlsx() -> None:
    folder = ROOT / "xlsx"
    folder.mkdir(parents=True)
    for number in (1, 2):
        workbook = Workbook()
        sheet = workbook.active
        sheet.title = "Clinical Data"
        sheet.append(["Field", "Synthetic Value", "Review"])
        for cell in sheet[1]:
            cell.font = Font(bold=True, color="FFFFFF")
            cell.fill = PatternFill("solid", fgColor="1F4E78")
        for line in phi_lines():
            key, value = line.split(":", 1)
            sheet.append([key, value.strip(), "Redaction validation"])
        sheet.append(["Calculated", "=COUNTA(B2:B10)", "Formula preservation"])
        if number == 2:
            second = workbook.create_sheet("Claims")
            second.append(["Procedure Code", "Diagnosis Code", "Payer ID"])
            second.append(["99213", "E11.9", "842610001"])
        workbook.save(folder / f"structured_workbook_{number}.xlsx")

    # Deliberate capability-probe fixture. Current XLSX processing handles cells,
    # not OCR inside drawing-layer images; the manifest calls this out explicitly.
    probe = Workbook()
    sheet = probe.active
    sheet["A1"] = f"Native cell email: {SYNTHETIC['email']}"
    image_path = folder / "_xlsx_image.png"
    text_image("Embedded spreadsheet image", phi_lines()[:4], (900, 500)).save(image_path)
    sheet.add_image(XlsxImage(str(image_path)), "A3")
    probe.save(folder / "capability_probe_embedded_image.xlsx")
    image_path.unlink()


def create_images() -> None:
    specifications = (("jpeg", "JPEG"), ("png", "PNG"), ("tiff", "TIFF"))
    for extension, image_format in specifications:
        folder = ROOT / extension
        folder.mkdir(parents=True)
        for number in (1, 2):
            image = text_image(f"Synthetic {image_format} Record {number}", phi_lines())
            image.save(folder / f"ocr_record_{number}.{extension if extension != 'jpeg' else 'jpg'}", format=image_format)


def create_dicom() -> None:
    folder = ROOT / "dicom"
    folder.mkdir(parents=True)
    for number in (1, 2):
        image = text_image(f"DICOM FRAME {number}", phi_lines()[:4], (1000, 700)).convert("L")
        array = np.asarray(image, dtype=np.uint8)
        sop_uid = generate_uid()
        meta = FileMetaDataset()
        meta.MediaStorageSOPClassUID = "1.2.840.10008.5.1.4.1.1.7"
        meta.MediaStorageSOPInstanceUID = sop_uid
        meta.TransferSyntaxUID = ExplicitVRLittleEndian
        dataset = FileDataset(str(folder / f"dicom_record_{number}.dcm"), {}, file_meta=meta, preamble=b"\0" * 128)
        dataset.SOPClassUID, dataset.SOPInstanceUID = meta.MediaStorageSOPClassUID, sop_uid
        dataset.StudyInstanceUID, dataset.SeriesInstanceUID = generate_uid(), generate_uid()
        dataset.PatientName, dataset.PatientID = SYNTHETIC["name"], SYNTHETIC["mrn"]
        dataset.PatientBirthDate = "19840115"
        dataset.Modality, dataset.StudyDescription = "OT", f"Synthetic study {SYNTHETIC['email']}"
        dataset.Rows, dataset.Columns = array.shape
        dataset.SamplesPerPixel, dataset.PhotometricInterpretation = 1, "MONOCHROME2"
        dataset.BitsAllocated = dataset.BitsStored = 8
        dataset.HighBit, dataset.PixelRepresentation = 7, 0
        dataset.PixelData = array.tobytes()
        dataset.save_as(dataset.filename, enforce_file_format=True)


def email_message(subject: str, *, attach_pdf: Path | None = None, inline_image: bool = False) -> EmailMessage:
    message = EmailMessage()
    message["From"], message["To"] = SYNTHETIC["email"], "reviewer@example.org"
    message["Subject"] = subject
    message.set_content("\n".join(phi_lines()))
    if inline_image:
        buffer = io.BytesIO()
        text_image("Inline scan", phi_lines()[:4], (900, 500)).save(buffer, format="PNG")
        message.add_attachment(buffer.getvalue(), maintype="image", subtype="png", filename="inline_scan.png")
    if attach_pdf:
        message.add_attachment(attach_pdf.read_bytes(), maintype="application", subtype="pdf", filename="clinical_attachment.pdf")
    return message


def create_email() -> None:
    pdf_attachment = ROOT / "pdf" / "patient_portal_mode.pdf"
    eml_folder = ROOT / "eml"
    eml_folder.mkdir(parents=True)
    (eml_folder / "message_with_pdf_attachment.eml").write_bytes(email_message("Synthetic referral", attach_pdf=pdf_attachment).as_bytes())
    (eml_folder / "message_with_inline_image.eml").write_bytes(email_message("Synthetic inline scan", inline_image=True).as_bytes())
    mbox_folder = ROOT / "mbox"
    mbox_folder.mkdir(parents=True)
    for number in (1, 2):
        path = mbox_folder / f"archive_{number}.mbox"
        archive = mailbox.mbox(path, create=True)
        archive.add(email_message(f"Synthetic archive message {number}", inline_image=number == 2))
        archive.add(email_message(f"Second synthetic message {number}", attach_pdf=pdf_attachment))
        archive.flush()
        archive.close()


def main() -> None:
    if ROOT.exists():
        shutil.rmtree(ROOT)
    ROOT.mkdir(parents=True)
    create_pdfs()
    create_docx()
    create_xlsx()
    create_images()
    create_dicom()
    create_email()
    print(f"Generated synthetic sample corpus at {ROOT}")


if __name__ == "__main__":
    main()
