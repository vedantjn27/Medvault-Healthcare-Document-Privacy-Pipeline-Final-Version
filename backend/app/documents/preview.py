"""Bounded, authenticated previews for every supported ingestion format."""

from __future__ import annotations

import base64
import io
import mailbox
import warnings
from email import policy
from email.message import Message
from email.parser import BytesParser
from pathlib import Path
from typing import Any

import fitz
import numpy as np
import pydicom
from docx import Document as DocxDocument
from openpyxl import load_workbook
from PIL import Image, ImageOps, ImageSequence, UnidentifiedImageError
from pydantic import BaseModel, Field

from app.documents.file_types import FileType


MAX_PREVIEW_PAGES = 10
MAX_PREVIEW_TEXT = 20_000
MAX_SHEETS = 20
MAX_SHEET_ROWS = 20
MAX_SHEET_COLUMNS = 20
THUMBNAIL_SIZE = (320, 320)
FULL_PAGE_MAX_SIDE = 1_800


class PreviewError(ValueError):
    pass


class PreviewResponse(BaseModel):
    file_type: FileType
    truncated: bool = False
    pages: list[dict[str, Any]] = Field(default_factory=list)
    text: str | None = None
    sheets: list[dict[str, Any]] = Field(default_factory=list)
    messages: list[dict[str, Any]] = Field(default_factory=list)
    metadata: dict[str, Any] = Field(default_factory=dict)


def _thumbnail_data_url(image: Image.Image) -> str:
    rendered = ImageOps.exif_transpose(image).convert("RGB")
    rendered.thumbnail(THUMBNAIL_SIZE)
    buffer = io.BytesIO()
    rendered.save(buffer, format="JPEG", quality=78, optimize=True)
    encoded = base64.b64encode(buffer.getvalue()).decode("ascii")
    return f"data:image/jpeg;base64,{encoded}"


def _full_png_bytes(image: Image.Image) -> bytes:
    """Render a readable, bounded PNG for an authenticated full-page preview."""

    rendered = ImageOps.exif_transpose(image).convert("RGB")
    rendered.thumbnail((FULL_PAGE_MAX_SIDE, FULL_PAGE_MAX_SIDE))
    buffer = io.BytesIO()
    rendered.save(buffer, format="PNG", optimize=True)
    return buffer.getvalue()


def _bounded_text(value: str, limit: int = MAX_PREVIEW_TEXT) -> tuple[str, bool]:
    return (value[:limit], len(value) > limit)


def _pdf_preview(path: Path) -> PreviewResponse:
    try:
        document = fitz.open(path)
    except (fitz.FileDataError, RuntimeError) as exc:
        raise PreviewError("PDF could not be opened") from exc
    try:
        if document.needs_pass:
            raise PreviewError("Password-protected PDFs are not supported")
        page_count = document.page_count
        pages: list[dict[str, Any]] = []
        for index in range(min(page_count, MAX_PREVIEW_PAGES)):
            page = document.load_page(index)
            text, text_truncated = _bounded_text(page.get_text("text"), 5_000)
            longest_side = max(page.rect.width, page.rect.height, 1.0)
            scale = min(1.0, 1200.0 / longest_side)
            matrix = fitz.Matrix(scale, scale)
            pixmap = page.get_pixmap(matrix=matrix, alpha=False)
            image = Image.open(io.BytesIO(pixmap.tobytes("png")))
            pages.append(
                {
                    "page_number": index + 1,
                    "text": text,
                    "text_truncated": text_truncated,
                    "thumbnail": _thumbnail_data_url(image),
                }
            )
        return PreviewResponse(
            file_type=FileType.PDF,
            pages=pages,
            truncated=page_count > MAX_PREVIEW_PAGES,
            metadata={"page_count": page_count},
        )
    finally:
        document.close()


def _docx_preview(path: Path) -> PreviewResponse:
    try:
        document = DocxDocument(path)
    except Exception as exc:
        raise PreviewError("Word document could not be opened") from exc
    parts: list[str] = []
    parts.extend(paragraph.text for paragraph in document.paragraphs if paragraph.text)
    for table in document.tables:
        for row in table.rows:
            parts.append("\t".join(cell.text for cell in row.cells))
    for section in document.sections:
        for container in (section.header, section.footer):
            parts.extend(paragraph.text for paragraph in container.paragraphs if paragraph.text)
    text, truncated = _bounded_text("\n".join(parts))
    return PreviewResponse(file_type=FileType.DOCX, text=text, truncated=truncated)


def _xlsx_preview(path: Path) -> PreviewResponse:
    try:
        workbook = load_workbook(path, read_only=True, data_only=False, keep_links=False)
    except Exception as exc:
        raise PreviewError("Excel workbook could not be opened") from exc
    try:
        sheets: list[dict[str, Any]] = []
        for worksheet in workbook.worksheets[:MAX_SHEETS]:
            rows = []
            for row in worksheet.iter_rows(
                min_row=1,
                max_row=min(worksheet.max_row, MAX_SHEET_ROWS),
                min_col=1,
                max_col=min(worksheet.max_column, MAX_SHEET_COLUMNS),
                values_only=True,
            ):
                rows.append([None if value is None else str(value) for value in row])
            sheets.append(
                {
                    "name": worksheet.title,
                    "rows": rows,
                    "row_count": worksheet.max_row,
                    "column_count": worksheet.max_column,
                    "truncated": (
                        worksheet.max_row > MAX_SHEET_ROWS
                        or worksheet.max_column > MAX_SHEET_COLUMNS
                    ),
                }
            )
        return PreviewResponse(
            file_type=FileType.XLSX,
            sheets=sheets,
            truncated=len(workbook.sheetnames) > MAX_SHEETS,
            metadata={"sheet_count": len(workbook.sheetnames)},
        )
    finally:
        workbook.close()


def _image_preview(path: Path, file_type: FileType) -> PreviewResponse:
    try:
        with warnings.catch_warnings():
            warnings.simplefilter("error", Image.DecompressionBombWarning)
            with Image.open(path) as image:
                frames = []
                for index, frame in enumerate(ImageSequence.Iterator(image)):
                    if index >= MAX_PREVIEW_PAGES:
                        break
                    frames.append(
                        {
                            "page_number": index + 1,
                            "thumbnail": _thumbnail_data_url(frame.copy()),
                        }
                    )
                frame_count = getattr(image, "n_frames", 1)
                return PreviewResponse(
                    file_type=file_type,
                    pages=frames,
                    truncated=frame_count > MAX_PREVIEW_PAGES,
                    metadata={"frame_count": frame_count, "width": image.width, "height": image.height},
                )
    except (
        UnidentifiedImageError,
        OSError,
        Image.DecompressionBombError,
        Image.DecompressionBombWarning,
    ) as exc:
        raise PreviewError("Image could not be opened safely") from exc


def _dicom_preview(path: Path) -> PreviewResponse:
    try:
        dataset = pydicom.dcmread(path)
    except Exception as exc:
        raise PreviewError("DICOM file could not be opened") from exc

    pages: list[dict[str, Any]] = []
    try:
        pixels = np.asarray(dataset.pixel_array)
        frames = pixels if pixels.ndim >= 3 and int(getattr(dataset, "NumberOfFrames", 1)) > 1 else [pixels]
        for index, frame in enumerate(frames[:MAX_PREVIEW_PAGES]):
            array = np.asarray(frame).astype(np.float32)
            minimum, maximum = float(array.min()), float(array.max())
            if maximum > minimum:
                array = (array - minimum) * (255.0 / (maximum - minimum))
            else:
                array = np.zeros_like(array)
            rendered = Image.fromarray(array.astype(np.uint8))
            if getattr(dataset, "PhotometricInterpretation", "") == "MONOCHROME1":
                rendered = ImageOps.invert(rendered)
            pages.append(
                {"page_number": index + 1, "thumbnail": _thumbnail_data_url(rendered)}
            )
    except Exception:
        pages = []

    frame_count = int(getattr(dataset, "NumberOfFrames", 1))
    return PreviewResponse(
        file_type=FileType.DICOM,
        pages=pages,
        truncated=frame_count > MAX_PREVIEW_PAGES,
        metadata={
            "frame_count": frame_count,
            "rows": int(getattr(dataset, "Rows", 0)),
            "columns": int(getattr(dataset, "Columns", 0)),
            "modality": str(getattr(dataset, "Modality", "")),
            "pixel_preview_available": bool(pages),
        },
    )


def _message_text(message: Message) -> str:
    parts: list[str] = []
    if message.is_multipart():
        for part in message.walk():
            if part.get_content_type() == "text/plain" and not part.get_filename():
                try:
                    parts.append(part.get_content())
                except (LookupError, UnicodeError):
                    continue
    elif message.get_content_type() == "text/plain":
        try:
            parts.append(message.get_content())
        except (LookupError, UnicodeError):
            pass
    return "\n".join(parts)


def _message_summary(message: Message) -> dict[str, Any]:
    body, truncated = _bounded_text(_message_text(message), 10_000)
    attachments = [part.get_filename() for part in message.walk() if part.get_filename()]
    return {
        "from": str(message.get("From", "")),
        "to": str(message.get("To", "")),
        "cc": str(message.get("Cc", "")),
        "subject": str(message.get("Subject", "")),
        "date": str(message.get("Date", "")),
        "body": body,
        "body_truncated": truncated,
        "attachments": attachments,
    }


def _eml_preview(path: Path) -> PreviewResponse:
    try:
        with path.open("rb") as source:
            message = BytesParser(policy=policy.default).parse(source)
    except Exception as exc:
        raise PreviewError("Email message could not be parsed") from exc
    return PreviewResponse(file_type=FileType.EML, messages=[_message_summary(message)])


def _mbox_preview(path: Path) -> PreviewResponse:
    archive = mailbox.mbox(path, create=False)
    messages: list[dict[str, Any]] = []
    try:
        for index, raw_message in enumerate(archive):
            if index >= MAX_PREVIEW_PAGES:
                break
            message = BytesParser(policy=policy.default).parsebytes(raw_message.as_bytes())
            messages.append(_message_summary(message))
        count = len(archive)
    except Exception as exc:
        raise PreviewError("Mailbox archive could not be parsed") from exc
    finally:
        archive.close()
    return PreviewResponse(
        file_type=FileType.MBOX,
        messages=messages,
        truncated=count > MAX_PREVIEW_PAGES,
        metadata={"message_count": count},
    )


def build_preview(path: Path, file_type: FileType) -> PreviewResponse:
    """Create a bounded preview using the format's real parser."""

    if file_type == FileType.PDF:
        return _pdf_preview(path)
    if file_type == FileType.DOCX:
        return _docx_preview(path)
    if file_type == FileType.XLSX:
        return _xlsx_preview(path)
    if file_type in {FileType.JPEG, FileType.PNG, FileType.TIFF}:
        return _image_preview(path, file_type)
    if file_type == FileType.DICOM:
        return _dicom_preview(path)
    if file_type == FileType.EML:
        return _eml_preview(path)
    if file_type == FileType.MBOX:
        return _mbox_preview(path)
    raise PreviewError("No preview handler exists for this file type")


def render_preview_page(path: Path, file_type: FileType, page_number: int) -> bytes:
    """Return one readable PNG page/frame without embedding it in the normal preview payload."""

    if page_number < 1 or page_number > MAX_PREVIEW_PAGES:
        raise PreviewError(f"Page must be between 1 and {MAX_PREVIEW_PAGES}")

    if file_type == FileType.PDF:
        try:
            document = fitz.open(path)
        except (fitz.FileDataError, RuntimeError) as exc:
            raise PreviewError("PDF could not be opened") from exc
        try:
            if document.needs_pass:
                raise PreviewError("Password-protected PDFs are not supported")
            if page_number > document.page_count:
                raise PreviewError("Page does not exist")
            page = document.load_page(page_number - 1)
            longest_side = max(page.rect.width, page.rect.height, 1.0)
            scale = min(2.4, FULL_PAGE_MAX_SIDE / longest_side)
            pixmap = page.get_pixmap(matrix=fitz.Matrix(scale, scale), alpha=False)
            return pixmap.tobytes("png")
        finally:
            document.close()

    if file_type in {FileType.JPEG, FileType.PNG, FileType.TIFF}:
        try:
            with Image.open(path) as image:
                frame_count = getattr(image, "n_frames", 1)
                if page_number > frame_count:
                    raise PreviewError("Frame does not exist")
                image.seek(page_number - 1)
                return _full_png_bytes(image.copy())
        except (UnidentifiedImageError, OSError) as exc:
            raise PreviewError("Image could not be opened safely") from exc

    if file_type == FileType.DICOM:
        try:
            dataset = pydicom.dcmread(path)
            pixels = np.asarray(dataset.pixel_array)
            frames = (
                pixels
                if pixels.ndim >= 3 and int(getattr(dataset, "NumberOfFrames", 1)) > 1
                else [pixels]
            )
            if page_number > len(frames):
                raise PreviewError("Frame does not exist")
            array = np.asarray(frames[page_number - 1]).astype(np.float32)
            minimum, maximum = float(array.min()), float(array.max())
            if maximum > minimum:
                array = (array - minimum) * (255.0 / (maximum - minimum))
            else:
                array = np.zeros_like(array)
            rendered = Image.fromarray(array.astype(np.uint8))
            if getattr(dataset, "PhotometricInterpretation", "") == "MONOCHROME1":
                rendered = ImageOps.invert(rendered)
            return _full_png_bytes(rendered)
        except PreviewError:
            raise
        except Exception as exc:
            raise PreviewError("DICOM preview could not be rendered") from exc

    raise PreviewError("Full-page viewing is available for PDF, image, TIFF, and DICOM files")
